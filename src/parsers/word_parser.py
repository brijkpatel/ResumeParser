"""Word document parser implementation."""

from pathlib import Path
from typing import Union

from docx import Document

from interfaces import FileParser
from exceptions import FileParsingError
from utils import logger


class WordParser(FileParser):
    """Concrete implementation of FileParser for Word documents.

    Uses python-docx library to extract text content from .docx files.
    Also attempts to handle .doc files when possible.
    """

    def __init__(self):
        """Initialize the Word parser."""
        self.supported_extensions = [".docx", ".doc"]
        logger.debug("WordParser initialized")

    def parse(self, file_path: str) -> str:
        """Parse a Word document and extract its text content.

        Args:
            file_path: Path to the Word document to parse

        Returns:
            Extracted text content from the document

        Raises:
            FileParsingError: If Word parsing fails
            FileNotFoundError: If file does not exist
        """
        self._validate_file_path(file_path)

        if not self.supports_format(file_path):
            raise FileParsingError("Unsupported file format.")

        logger.info("Parsing Word document: %s", file_path)

        # Handle .doc files (legacy format)
        if self._get_file_extension(file_path) == ".doc":
            return self._parse_doc_file(file_path)

        # Handle .docx files
        try:
            document = Document(file_path)
            text_content: list[str] = []

            # Extract text from paragraphs
            for paragraph in document.paragraphs:
                text = paragraph.text.strip()
                if text:
                    text_content.append(text)

            # Extract text from tables
            for table in document.tables:
                for row in table.rows:
                    row_text: list[str] = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        text_content.append(" | ".join(row_text))

            if not text_content:
                raise FileParsingError("No text content found in document.")

            # Join all content with newlines
            full_text = "\n".join(text_content)

            logger.info(
                "Successfully parsed Word document: %s (%d characters)",
                file_path,
                len(full_text),
            )

            return full_text

        except Exception as e:
            raise FileParsingError(
                "Error parsing Word document.", original_exception=e
            ) from e

    def _parse_doc_file(self, file_path: str) -> str:
        """Parse legacy .doc files.

        For .doc files, we provide a fallback approach since python-docx
        doesn't support the legacy format natively.

        Args:
            file_path: Path to the .doc file

        Returns:
            Extracted text content

        Raises:
            FileParsingError: If parsing fails
        """
        logger.warning("Legacy .doc format detected: %s", file_path)

        try:
            # Try to use python-docx anyway (sometimes works with newer .doc files)
            document = Document(file_path)
            text_content: list[str] = []

            for paragraph in document.paragraphs:
                text = paragraph.text.strip()
                if text:
                    text_content.append(text)

            if text_content:
                full_text = "\n".join(text_content)
                logger.info("Successfully parsed .doc file using docx: %s", file_path)
                return full_text

        except Exception as e:
            raise FileParsingError(
                "Error parsing Word document.", original_exception=e
            ) from e

        # If python-docx fails, suggest alternative approaches
        raise FileParsingError(
            "Cannot parse legacy .doc file. Please convert to .docx format or use a tool like LibreOffice to convert the file."
        )

    def supports_format(self, file_path: Union[str, Path]) -> bool:
        """Check if this parser supports the given file format.

        Args:
            file_path: Path to the file to check

        Returns:
            True if the file is a Word document, False otherwise
        """
        extension = self._get_file_extension(file_path)
        return extension in self.supported_extensions
